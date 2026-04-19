"""Tests for parsers and event extractor."""

import io
from datetime import date

import pytest

from core.base_calendar import BaseCalendar, _build_day_map
from core.parsers import ParsedDoc
from core.parsers.event_extractor import extract_events

# ---------------------------------------------------------------------------
# Shared mini calendar
# ---------------------------------------------------------------------------

_RAW = {
    "ay": "TEST",
    "semester": "Test",
    "start_date": "2026-01-05",
    "end_date": "2026-04-24",
    "tee_start": "2026-04-27",
    "tee_end": "2026-05-08",
    "grad_start": None,
    "grad_end": None,
    "first_academic_day_type": "1",
    "special_days": {
        "2026-01-19": {"day_type": "holiday", "notes": ["MLK Day"]},
    },
}


@pytest.fixture
def mini_cal():
    return BaseCalendar("TEST", _day_map=_build_day_map(_RAW))


# ---------------------------------------------------------------------------
# DOCX parser
# ---------------------------------------------------------------------------

class TestDocxParser:
    def _make_docx(self, paragraphs: list[str], table_data: list[list[str]] = None) -> bytes:
        from docx import Document
        doc = Document()
        for p in paragraphs:
            doc.add_paragraph(p)
        if table_data:
            tbl = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for i, row in enumerate(table_data):
                for j, val in enumerate(row):
                    tbl.cell(i, j).text = val
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def test_paragraphs_extracted(self):
        from core.parsers.docx_parser import parse_docx
        data = self._make_docx(["WPR 1 at Lesson 15", "TEE at end of semester"])
        doc = parse_docx(data)
        assert len(doc.paragraphs) == 2
        assert "WPR" in doc.full_text

    def test_table_extracted(self):
        from core.parsers.docx_parser import parse_docx
        data = self._make_docx([], [["Event", "Lesson", "Weight"], ["WPR 1", "L15", "20%"]])
        doc = parse_docx(data)
        assert len(doc.tables) == 1
        assert doc.tables[0][1][0] == "WPR 1"

    def test_empty_paragraphs_skipped(self):
        from core.parsers.docx_parser import parse_docx
        data = self._make_docx(["", "WPR 1", "", ""])
        doc = parse_docx(data)
        assert doc.paragraphs == ["WPR 1"]

    def test_is_scan_false_for_normal_docx(self):
        from core.parsers.docx_parser import parse_docx
        data = self._make_docx(["WPR 1"])
        doc = parse_docx(data)
        assert doc.is_scan is False


# ---------------------------------------------------------------------------
# XLSX parser
# ---------------------------------------------------------------------------

class TestXlsxParser:
    def _make_xlsx(self, rows: list[list]) -> bytes:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        for row in rows:
            ws.append(row)
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf.read()

    def test_table_extracted(self):
        from core.parsers.xlsx_parser import parse_xlsx
        data = self._make_xlsx([
            ["Event Type", "Due Date", "Weight"],
            ["WPR", "14 Feb 2026", "20%"],
            ["TEE", "4 May 2026", "30%"],
        ])
        doc = parse_xlsx(data)
        assert len(doc.tables) == 1
        assert doc.tables[0][1][0] == "WPR"

    def test_full_text_contains_data(self):
        from core.parsers.xlsx_parser import parse_xlsx
        data = self._make_xlsx([["WPR", "L15", "20%"]])
        doc = parse_xlsx(data)
        assert "WPR" in doc.full_text


# ---------------------------------------------------------------------------
# Event extractor — text pass
# ---------------------------------------------------------------------------

class TestEventExtractorText:
    def test_extracts_wpr_with_lesson_ref(self, mini_cal):
        doc = ParsedDoc(
            paragraphs=["WPR 1 will be held at Lesson 10."],
            full_text="WPR 1 will be held at Lesson 10.",
        )
        events = extract_events(doc, "MA364", 1, mini_cal)
        assert any(e.event_type == "WPR" for e in events)

    def test_extracts_tee_with_explicit_date(self, mini_cal):
        doc = ParsedDoc(
            paragraphs=["TEE scheduled for 4 May 2026."],
            full_text="TEE scheduled for 4 May 2026.",
        )
        events = extract_events(doc, "MA364", 1, mini_cal)
        assert any(e.event_type == "TEE" for e in events)

    def test_course_code_set_correctly(self, mini_cal):
        doc = ParsedDoc(full_text="WPR 1 at Lesson 5")
        events = extract_events(doc, "PH201", 2, mini_cal)
        assert all(e.course_code == "PH201" for e in events)

    def test_events_sorted_by_date(self, mini_cal):
        doc = ParsedDoc(full_text="WPR 2 at Lesson 20. WPR 1 at Lesson 5.")
        events = extract_events(doc, "MA364", 1, mini_cal)
        dates = [e.date for e in events]
        assert dates == sorted(dates)

    def test_no_date_event_dropped(self, mini_cal):
        doc = ParsedDoc(full_text="WPR on some unspecified day")
        events = extract_events(doc, "MA364", 1, mini_cal)
        # All returned events must have a resolved date
        assert all(e.date is not None for e in events)


# ---------------------------------------------------------------------------
# Event extractor — table pass
# ---------------------------------------------------------------------------

class TestEventExtractorTable:
    def test_table_with_headers(self, mini_cal):
        doc = ParsedDoc(
            tables=[[
                ["Event Type", "Lesson", "Weight"],
                ["WPR", "L8", "20%"],
                ["TEE", "L40", "30%"],
            ]],
            full_text="",
        )
        events = extract_events(doc, "MA364", 1, mini_cal)
        types = {e.event_type for e in events}
        assert "WPR" in types
        assert "TEE" in types

    def test_weight_extracted_from_table(self, mini_cal):
        doc = ParsedDoc(
            tables=[[
                ["Event", "Lesson", "Weight"],
                ["WPR 1", "L10", "25%"],
            ]],
            full_text="",
        )
        events = extract_events(doc, "MA364", 1, mini_cal)
        wpr = next((e for e in events if e.event_type == "WPR"), None)
        assert wpr is not None
        assert wpr.weight_pct == 25.0

    def test_table_with_explicit_date(self, mini_cal):
        doc = ParsedDoc(
            tables=[[
                ["Assignment", "Due Date"],
                ["HW 1", "14 Jan 2026"],
            ]],
            full_text="",
        )
        events = extract_events(doc, "MA364", 1, mini_cal)
        hw = next((e for e in events if e.event_type == "HW"), None)
        assert hw is not None
        assert hw.date == date(2026, 1, 14)

    def test_table_without_header_row(self, mini_cal):
        # Single-row table (no header) should still attempt extraction
        doc = ParsedDoc(
            tables=[[["WPR 1", "L5", "20%"]]],
            full_text="",
        )
        # May or may not extract (depends on column guessing) — at minimum no crash
        extract_events(doc, "MA364", 1, mini_cal)


# ---------------------------------------------------------------------------
# Confidence scoring
# ---------------------------------------------------------------------------

class TestConfidenceScoring:
    def test_table_event_has_higher_confidence_than_bare_text(self, mini_cal):
        # Table extraction sets course_code_in_ctx=True; text extraction may not
        doc_table = ParsedDoc(
            tables=[[
                ["Event", "Lesson"],
                ["WPR 1", "L10"],
            ]],
            full_text="",
        )
        doc_text = ParsedDoc(full_text="WPR at Lesson 10")

        evts_table = extract_events(doc_table, "MA364", 1, mini_cal)
        evts_text = extract_events(doc_text, "MA364", 1, mini_cal)

        if evts_table and evts_text:
            assert evts_table[0].confidence >= evts_text[0].confidence

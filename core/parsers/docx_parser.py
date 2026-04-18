"""DOCX syllabus parser using python-docx."""

import io
from pathlib import Path
from typing import Union

from docx import Document

from core.parsers import ParsedDoc


def parse_docx(source: Union[Path, bytes, io.BytesIO]) -> ParsedDoc:
    """Parse a DOCX file and return structured content."""
    if isinstance(source, bytes):
        source = io.BytesIO(source)
    elif isinstance(source, Path):
        source = source.open("rb")

    doc = Document(source)
    result = ParsedDoc()
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        result.paragraphs.append(text)
        lines.append(text)
        style = para.style.name or ""
        if style.lower().startswith("heading"):
            try:
                level = int(style.split()[-1])
            except (ValueError, IndexError):
                level = 1
            result.headings.append((level, text))

    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            seen: set[int] = set()
            cells: list[str] = []
            for cell in row.cells:
                cid = id(cell._element)
                cells.append("" if cid in seen else cell.text.strip())
                seen.add(cid)
            rows.append(cells)
            lines.append("\t".join(c for c in cells if c))
        result.tables.append(rows)

    result.full_text = "\n".join(lines)
    return result
